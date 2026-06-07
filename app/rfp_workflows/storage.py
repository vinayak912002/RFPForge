# DB operations
import csv
import io
import json
import re
from pathlib import Path
from typing import Any

from app.knowledge_engine.prompting import build_question_extraction_prompt


MAX_CHARS_PER_LLM_CHUNK = 9000
CHUNK_OVERLAP_CHARS = 700
MAX_LLM_CHUNKS = 80

TEXT_EXTENSIONS = {".txt", ".md", ".rtf", ".log"}
CSV_EXTENSIONS = {".csv", ".tsv"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm"}

QUESTION_STARTERS = (
    "describe",
    "provide",
    "explain",
    "confirm",
    "submit",
    "include",
    "identify",
    "list",
    "detail",
    "outline",
    "summarize",
    "demonstrate",
    "specify",
    "state",
    "attach",
    "complete",
    "clarify",
    "define",
    "document",
    "share",
    "supply",
    "answer",
    "respond",
    "indicate",
)

RESPONSE_VERBS = (
    "provide",
    "describe",
    "explain",
    "confirm",
    "submit",
    "include",
    "identify",
    "list",
    "specify",
    "state",
    "complete",
    "clarify",
    "document",
    "demonstrate",
    "outline",
)

ACTOR_WORDS = (
    "vendor",
    "bidder",
    "proposer",
    "respondent",
    "supplier",
    "contractor",
    "applicant",
    "customer",
    "client",
    "team",
    "owner",
)

REQUIREMENT_WORDS = (
    "must",
    "shall",
    "should",
    "will",
    "is required to",
    "are required to",
    "needs to",
    "need to",
)

IMPLIED_NEED_WORDS = (
    "tbd",
    "to be determined",
    "to be decided",
    "unknown",
    "unclear",
    "not specified",
    "not defined",
    "missing",
    "pending",
    "needs confirmation",
    "requires confirmation",
    "requires clarification",
)

NOISE_PATTERNS = (
    r"^page \d+(\s+of\s+\d+)?$",
    r"^confidential$",
    r"^proprietary$",
    r"^table of contents$",
    r"^copyright\b",
)


def normalize_question(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\.{2,}$", "", text).strip()
    text = re.sub(r"^[\-\*\u2022\u25e6\u00b7\d]+[\.\)\:\-\s]+", "", text).strip()
    text = re.sub(r"^(q|question|requirement|item)\s*\d*[\.\:\-\)]\s*", "", text, flags=re.I)
    text = re.sub(r"\s+([,.;:?!])", r"\1", text)
    return text.strip(" \t\r\n;")


def is_noise(text: str) -> bool:
    normalized = normalize_question(text)
    lower = normalized.lower()

    if len(normalized) < 12:
        return True

    if len(normalized.split()) <= 3 and not normalized.endswith("?"):
        return True

    if any(re.match(pattern, lower) for pattern in NOISE_PATTERNS):
        return True

    return False


def dedupe_questions(questions: list[str]) -> list[str]:
    seen = set()
    unique = []

    for question in questions:
        normalized = normalize_question(question)
        if is_noise(normalized):
            continue

        key = re.sub(r"[^a-z0-9]+", " ", normalized.lower()).strip()
        if key not in seen:
            seen.add(key)
            unique.append(normalized)

    return unique


def chunk_text(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text) and len(chunks) < MAX_LLM_CHUNKS:
        end = min(start + MAX_CHARS_PER_LLM_CHUNK, len(text))

        if end < len(text):
            paragraph_break = text.rfind("\n\n", start, end)
            line_break = text.rfind("\n", start, end)
            best_break = max(paragraph_break, line_break)

            if best_break > start + (MAX_CHARS_PER_LLM_CHUNK // 2):
                end = best_break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(end - CHUNK_OVERLAP_CHARS, 0)

    return chunks


def extract_json_payload(response: str) -> dict[str, Any] | list[Any] | None:
    response = response.strip()

    if response.startswith("```"):
        response = re.sub(r"^```(?:json)?", "", response, flags=re.I).strip()
        response = re.sub(r"```$", "", response).strip()

    for candidate in (response,):
        try:
            data = json.loads(candidate)
            if isinstance(data, (dict, list)):
                return data
        except json.JSONDecodeError:
            pass

    json_match = re.search(r"(\{.*\}|\[.*\])", response, re.DOTALL)
    if not json_match:
        return None

    try:
        data = json.loads(json_match.group(0))
    except json.JSONDecodeError:
        return None

    return data if isinstance(data, (dict, list)) else None


def questions_from_json_payload(payload: dict[str, Any] | list[Any] | None) -> list[str]:
    if payload is None:
        return []

    if isinstance(payload, list):
        raw_items = payload
    else:
        raw_items = (
            payload.get("questions")
            or payload.get("items")
            or payload.get("requirements")
            or []
        )

    if not isinstance(raw_items, list):
        return []

    questions = []
    for item in raw_items:
        if isinstance(item, str):
            questions.append(item)
            continue

        if isinstance(item, dict):
            value = (
                item.get("question")
                or item.get("text")
                or item.get("requirement")
                or item.get("request")
            )
            if isinstance(value, str):
                questions.append(value)

    return questions


def extract_questions_from_plain_llm_text(response: str) -> list[str]:
    questions = []

    for line in response.splitlines():
        line = normalize_question(line)
        if not line:
            continue

        if line.endswith("..."):
            line = line[:-3].strip()

        if not line:
            continue

        if line.endswith("?") or line.lower().startswith(QUESTION_STARTERS):
            questions.append(line)

    return questions


def extract_questions_with_llm(text: str, llm_service) -> list[str]:
    questions = []

    for chunk in chunk_text(text):
        prompt = build_question_extraction_prompt(chunk)

        try:
            response = llm_service.generate(prompt)
        except Exception:
            continue

        payload = extract_json_payload(response)
        parsed_questions = questions_from_json_payload(payload)

        if parsed_questions:
            questions.extend(parsed_questions)
        else:
            questions.extend(extract_questions_from_plain_llm_text(response))

    return dedupe_questions(questions)


def extract_table_like_questions(text: str) -> list[str]:
    questions = []

    for line in text.splitlines():
        if "|" not in line and "\t" not in line:
            continue

        cells = [
            normalize_question(cell)
            for cell in re.split(r"\||\t", line)
            if normalize_question(cell)
        ]

        if len(cells) < 2:
            continue

        row_text = " - ".join(cells)
        lower = row_text.lower()

        if any(word in lower for word in RESPONSE_VERBS + IMPLIED_NEED_WORDS):
            questions.append(row_text)

    return questions


def extract_questions_rule_based(text: str) -> list[str]:
    questions = []
    candidates = []

    candidates.extend(text.splitlines())
    candidates.extend(re.split(r"(?<=[\?\.!])\s+", re.sub(r"\s+", " ", text)))
    questions.extend(extract_table_like_questions(text))

    for candidate in candidates:
        line = normalize_question(candidate)
        if not line:
            continue

        if "|" in line or "\t" in line:
            continue

        lower = line.lower()

        if line.endswith("?") and len(line) > 10:
            questions.append(line)
            continue

        if lower.startswith(QUESTION_STARTERS) and len(line) > 16:
            questions.append(line)
            continue

        actor_requirement = (
            rf"^({'|'.join(ACTOR_WORDS)}) "
            rf"({'|'.join(REQUIREMENT_WORDS)})\b"
        )
        if re.match(actor_requirement, lower):
            questions.append(line)
            continue

        if any(f" {verb} " in f" {lower} " for verb in RESPONSE_VERBS) and any(
            word in lower for word in REQUIREMENT_WORDS
        ):
            questions.append(line)
            continue

        if any(word in lower for word in IMPLIED_NEED_WORDS):
            questions.append(f"Clarify: {line}")
            continue

    return dedupe_questions(questions)


def read_upload_bytes(file) -> bytes:
    file.file.seek(0)
    data = file.file.read()
    file.file.seek(0)
    return data


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode("utf-8", errors="ignore")


def extract_text_from_pdf(file) -> str:
    from pypdf import PdfReader

    file.file.seek(0)
    reader = PdfReader(file.file)
    pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()
        if page_text:
            pages.append(f"[Page {page_number}]\n{page_text}")

    file.file.seek(0)
    return "\n\n".join(pages)


def extract_text_from_docx(file) -> str:
    from docx import Document

    file.file.seek(0)
    doc = Document(file.file)
    blocks = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    file.file.seek(0)
    return "\n".join(blocks)


def extract_text_from_spreadsheet(file) -> str:
    from openpyxl import load_workbook

    file.file.seek(0)
    workbook = load_workbook(file.file, read_only=True, data_only=True)
    rows = []

    for sheet in workbook.worksheets:
        rows.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))

    workbook.close()
    file.file.seek(0)
    return "\n".join(rows)


def extract_text_from_csv(file, delimiter: str = ",") -> str:
    data = read_upload_bytes(file)
    text = decode_text(data)
    rows = []

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    for row in reader:
        cells = [cell.strip() for cell in row if cell.strip()]
        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows)


def extract_text_from_file(file) -> str:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file)

    if extension == ".docx":
        return extract_text_from_docx(file)

    if extension in SPREADSHEET_EXTENSIONS:
        return extract_text_from_spreadsheet(file)

    if extension in CSV_EXTENSIONS:
        delimiter = "\t" if extension == ".tsv" else ","
        return extract_text_from_csv(file, delimiter=delimiter)

    if extension in TEXT_EXTENSIONS:
        return decode_text(read_upload_bytes(file))

    return decode_text(read_upload_bytes(file))


def extract_questions(text: str):
    return extract_questions_rule_based(text)


def parse_file(file, llm_service=None):
    text = extract_text_from_file(file)
    if not text.strip():
        return []

    if llm_service:
        questions = extract_questions_with_llm(text, llm_service)
        if questions:
            return questions

    return extract_questions_rule_based(text)
