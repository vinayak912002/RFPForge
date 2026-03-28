# DOCX / PDF export
from docx import Document
import pandas as pd
from app.rfp_workflows.models import Draft


def export_to_word(db, rfp_id: str):

    drafts = db.query(Draft).filter(
        Draft.rfp_id == rfp_id,
        Draft.status == "final"
    ).all()

    if not drafts:
        return None

    doc = Document()
    doc.add_heading("RFP Response", 0)

    for draft in drafts:
        doc.add_heading(f"Question ID: {draft.question_id}", level=2)
        doc.add_paragraph(draft.answer_text)

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix - Sources", level=1)
    doc.add_paragraph("Generated with RFP Tool")

    file_path = f"{rfp_id}_response.docx"
    doc.save(file_path)

    return file_path


def export_to_excel(db, rfp_id: str):

    drafts = db.query(Draft).filter(
        Draft.rfp_id == rfp_id,
        Draft.status == "final"
    ).all()

    if not drafts:
        return None

    data = []

    for d in drafts:
        data.append({
            "Question ID": d.question_id,
            "Answer": d.answer_text
        })

    df = pd.DataFrame(data)

    file_path = f"{rfp_id}_response.xlsx"
    df.to_excel(file_path, index=False)

    return file_path